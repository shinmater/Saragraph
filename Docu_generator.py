# -*- coding: utf-8 -*-
"""
Created on Thu Oct  7 19:20:56 2021

@author: Yongjin

Reference: https://besixdouze.net/40

"""

## Installation needed
## pip install python-docx



from docx import Document

document = Document()
document.save("test.docx")

# document.add_heading('이 문서의 타이틀! lvl = 0', level = 0)
# document.add_heading('이 문서의 타이틀! lvl = 1', level = 1)
# document.add_heading('이 문서의 타이틀! lvl = 2', level = 2)
# document.add_heading('이 문서의 타이틀! lvl = 3', level = 3)
# document.add_heading('이 문서의 타이틀! lvl = 4', level = 4)
# document.add_heading('이 문서의 타이틀! lvl = 5', level = 5)

from docx.shared import Inches 
document.add_picture('Logo.jpg', width=Inches(2.0))


from docx.shared import RGBColor, Pt

paragraph = document.add_paragraph("문단을 추가합시다. 이게 1번째 문단입니다. \n 줄바꿈이 되는지 해봅니다.")
paragraph = document.add_paragraph("문단을 추가합시다. 이게 2번째 문단입니다. 일부러 변수명을 똑같이 해봤습니다.") 
insert_paragraph = paragraph.insert_paragraph_before("3번째로 생성한 문단을 삽입합니다. 2번째 문단 앞에 넣습니다.") 
paragraph = document.add_paragraph("문단을 추가합시다. 이게 4번째 문단입니다.\n")
paragraph.add_run(" 4번째 문단에 글을 덧붙여 봅니다.")
paragraph = document.add_paragraph() # 5번째 빈 문단 생성 paragraph.add_run("5번째 문단. 이 문장이 Bold가 될까 해봅니다.\n").bold = True 
paragraph.add_run(" 5번째 문단. 이 문장의 글자색 바꿔봅니다. 귀찮아서 안 쓸 것 같습니다.\n").font.color.rgb = RGBColor(255, 0, 0) 
run = paragraph.add_run(" 5번째 문단. 이 문장은 Bold 적용, 글씨색 변경, 이텔릭, 글씨체, 글씨크기 변경을 다 적용해봅니다.\n") 
run.bold = True 
run.font.color.rgb = RGBColor(0, 0, 255) 
run.italic = True 
run.font.name = "궁서체" 
run.font.size = Pt(20) 
paragraph.add_run(" 정말 귀찮습니다. 그냥 미리 서식 파일 만들어 놓고 씁시다 (아직 5번째 문단임)\n")

run = paragraph.add_run(" This is test for Sarah. It seems that font change works well for english\n") 
run.bold = True 
run.font.color.rgb = RGBColor(0, 100, 255) 
run.italic = True 
run.font.name = "Times New Roman" 
run.font.size = Pt(20) 



table = document.add_table(rows = 2, cols = 3, style = "Table Grid") 

cell = table.cell(0,1) # 0번째 행, 1번째 열의 셀을 지정 
cell.text = "0,1" # 해당 셀에 "0,1"값 대입 
table.rows[1].cells[0].text = "1,0" # 1번째 행, 0번째 열에 "1,0"대입 
table.columns[2].cells[0].text = "0,2" # 0번째 행, 2번째 열에 "0,2"대입 
row = table.add_row() # 표에 행 추가 
#row.height = Inches(2.0)
row.cells[1].text = "2,1" # 추가된 행의 1번째 열에 "2,1"대입
row.cells[2].text = 'TEST'

para=row.cells[2].add_paragraph()
run = para.add_run('test first')
#table.style.font.name="Times New Roman"

row.table.style.font.name="Times New Roman"

# # 문서 내 모든 표의 모든 셀을 출력 
# for table in document.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             print(cell.text)
            
# # 문서 내 모든 문단 출력 
# for p in document.paragraphs:
#     print(p.text)

document.save("test.docx")

from docx2pdf import convert
convert("test.docx", "test3.pdf")


import sys
import os
import comtypes.client

wdFormatPDF = 17

in_file = "test.docx"
out_file = "test.pdf"


word = comtypes.client.CreateObject('Word.Application')
doc = word.Documents.Open(in_file)
doc.SaveAs(out_file, FileFormat=wdFormatPDF)
doc.Close()
word.Quit()